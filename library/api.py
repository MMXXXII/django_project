from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.db.models import Count
from django.http import HttpResponse
import io
from django.contrib.auth.models import User
from openpyxl import Workbook
from docx import Document
from library.models import Library, Book, Genre, Member, Loan, UserProfile
from library.serializers import (LibrarySerializer, BookSerializer, GenreSerializer,MemberSerializer, LoanSerializer, UserSerializer)


import io
from openpyxl import Workbook
from docx import Document
from django.http import HttpResponse
from rest_framework.response import Response


class BaseExportMixin:
    def export_queryset(self, queryset, columns, filename_base):
        file_type = self.request.query_params.get('type', 'excel')

        if file_type == 'excel':
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = filename_base
            
            sheet.append(columns)
            
            for data_row in queryset:
                row_data = []
                for col in columns:
                    value = data_row.get(col, '')
                    row_data.append(value)
                sheet.append(row_data)

            excel_file = io.BytesIO()
            workbook.save(excel_file)
            excel_file.seek(0)

            file_response = HttpResponse(
                excel_file,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            file_response['Content-Disposition'] = f'attachment; filename="{filename_base}.xlsx"'
            return file_response

        elif file_type == 'word':
            document = Document()
            document.add_heading(filename_base, 0)

            for data_row in queryset:
                row_values = []
                for col in columns:
                    value = data_row.get(col, '')
                    row_values.append(str(value))
                
                text_line = ' | '.join(row_values)
                document.add_paragraph(text_line)

            word_file = io.BytesIO()
            document.save(word_file)
            word_file.seek(0)

            file_response = HttpResponse(
                word_file,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            file_response['Content-Disposition'] = f'attachment; filename="{filename_base}.docx"'
            return file_response

        else:
            return Response({"error": "Unknown file type"}, status=400)


class GenreViewSet(viewsets.ModelViewSet, BaseExportMixin):
    serializer_class = GenreSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Genre.objects.all().order_by('name')

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    @action(detail=False, methods=['get'])
    def stats(self, request):
        queryset = self.get_queryset()
        total = queryset.count()
        
        top_genre = Genre.objects.annotate(book_count=Count('book')).order_by('-book_count').first()
        top_name = top_genre.name if top_genre else None

        return Response({
            'count': total,
            'top': top_name
        })

    @action(detail=False, methods=['get'])
    def export(self, request):
        queryset = self.get_queryset()
        
        if request.user.is_superuser:
            data_list = []
            for genre in queryset:
                user_name = genre.user.username if genre.user else ''
                data_list.append({
                    'ID': genre.id,
                    'Name': genre.name,
                    'User': user_name
                })
            columns = ['ID', 'Name', 'User']
        else:
            data_list = []
            for genre in queryset:
                data_list.append({
                    'ID': genre.id,
                    'Name': genre.name
                })
            columns = ['ID', 'Name']
        
        return self.export_queryset(data_list, columns, 'Genres')

class LibraryViewSet(viewsets.ModelViewSet, BaseExportMixin):
    serializer_class = LibrarySerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Library.objects.all().order_by('name')

    def perform_create(self, serializer):
        if not self.request.user.is_superuser:
            raise PermissionError("Only admins can add libraries.")
        serializer.save(user=self.request.user)

    def perform_update(self, serializer):
        if not self.request.user.is_superuser:
            raise PermissionError("Only admins can edit libraries.")
        serializer.save(user=self.request.user)

    def perform_destroy(self, instance):
        if not self.request.user.is_superuser:
            raise PermissionError("Only admins can delete libraries.")
        instance.delete()

    @action(detail=False, methods=['get'])
    def stats(self, request):
        queryset = self.get_queryset()
        total = queryset.count()
        
        top_library = Library.objects.annotate(loan_count=Count('book__loan')).order_by('-loan_count').first()
        top_name = top_library.name if top_library else None

        return Response({
            'count': total,
            'top': top_name
        })

    @action(detail=False, methods=['get'])
    def export(self, request):
        queryset = self.get_queryset()
        
        if request.user.is_superuser:
            data_list = []
            for library in queryset:
                user_name = library.user.username if library.user else ''
                data_list.append({
                    'ID': library.id,
                    'Name': library.name,
                    'User': user_name
                })
            columns = ['ID', 'Name', 'User']
        else:
            data_list = []
            for library in queryset:
                data_list.append({
                    'ID': library.id,
                    'Name': library.name
                })
            columns = ['ID', 'Name']
        
        return self.export_queryset(data_list, columns, 'Libraries')


class BookViewSet(viewsets.ModelViewSet, BaseExportMixin):
    serializer_class = BookSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Book.objects.all()

    @action(detail=False, methods=['get'])
    def stats(self, request):
        queryset = self.get_queryset()
        total = queryset.count()

        most_borrowed = Loan.objects.values('book__id', 'book__title').annotate(borrow_count=Count('book')).order_by('-borrow_count').first()

        if most_borrowed:
            most_borrowed_book = {
                'id': most_borrowed['book__id'],
                'title': most_borrowed['book__title'],
                'borrow_count': most_borrowed['borrow_count']
            }
        else:
            most_borrowed_book = None

        return Response({
            'count': total,
            'most_borrowed': most_borrowed_book
        })

    @action(detail=False, methods=['get'])
    def export(self, request):
        if not request.user.is_superuser:
            return Response(
                {"error": "Only administrators can export books"}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        queryset = self.get_queryset()
        data_list = []
        
        for book in queryset:
            genre_name = book.genre.name if book.genre else ''
            library_name = book.library.name if book.library else ''
            status_text = 'Available' if book.is_available else 'Borrowed'
            
            data_list.append({
                'ID': book.id,
                'Title': book.title,
                'Genre': genre_name,
                'Library': library_name,
                'Status': status_text
            })
        
        columns = ['ID', 'Title', 'Genre', 'Library', 'Status']
        
        return self.export_queryset(data_list, columns, 'Books')
    

class LoanViewSet(viewsets.ModelViewSet, BaseExportMixin):
    serializer_class = LoanSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        queryset = Loan.objects.select_related('book', 'member', 'user')
        
        if self.request.user.is_superuser:
            return queryset
        
        return queryset.filter(member__user=self.request.user)

    @action(detail=True, methods=['post'], url_path='return')
    def return_book(self, request, pk=None):
        loan = self.get_object()
        
        if loan.return_date:
            return Response(
                {'detail': 'Книга уже возвращена.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        loan.return_date = date.today()
        loan.save()
        
        self.update_loan_stats()

        serializer = self.get_serializer(loan)
        return Response({
            'detail': 'Книга успешно возвращена.',
            'loan': serializer.data
        })

    def update_loan_stats(self):
        top_reader = Loan.objects.filter(return_date__isnull=True).values('member__id', 'member__first_name').annotate(loan_count=Count('id')).order_by('-loan_count').first()

        top_reader_data = {
            'name': top_reader['member__first_name'] if top_reader else None,
            'loan_count': top_reader['loan_count'] if top_reader else 0
        }

        loan_count = Loan.objects.filter(return_date__isnull=False).count()

        return Response({
            'count': loan_count,
            'topReader': top_reader_data
        })

    @action(detail=False, methods=['get'])
    def stats(self, request):
        queryset = self.get_queryset()
        total = queryset.count()

        top_reader = queryset.values('member__id', 'member__first_name').annotate(loan_count=Count('id')).order_by('-loan_count').first()
        
        if top_reader:
            top_reader_data = {
                'name': top_reader['member__first_name'],
                'loan_count': top_reader['loan_count']
            }
        else:
            top_reader_data = {
                'name': None,
                'loan_count': 0
            }

        return Response({
            'count': total,
            'topReader': top_reader_data
        })

    @action(detail=False, methods=['get'])
    def export(self, request):
        if not request.user.is_superuser:
            return Response({"error": "Permission denied"}, status=403)

        queryset = self.get_queryset()
        data_list = []
        
        for loan in queryset:
            book_title = loan.book.title if loan.book else ''
            member_name = loan.member.first_name if loan.member else ''
            user_name = loan.user.username if loan.user else ''
            return_date = loan.return_date if loan.return_date else 'Not returned'
            
            data_list.append({
                'ID': loan.id,
                'Book': book_title,
                'Member': member_name,
                'User': user_name,
                'Loan Date': loan.loan_date,
                'Return Date': return_date
            })
        
        columns = ['ID', 'Book', 'Member', 'User', 'Loan Date', 'Return Date']
        
        return self.export_queryset(data_list, columns, 'Loans')


class MemberViewSet(viewsets.ModelViewSet, BaseExportMixin):
    serializer_class = UserSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        if self.request.user.is_superuser:
            return User.objects.all()
        return User.objects.filter(id=self.request.user.id)

    def create(self, request, *args, **kwargs):
        username = request.data.get('username')
        email = request.data.get('email', '')
        password = request.data.get('password')
        is_superuser = request.data.get('is_superuser', False)
        is_staff = request.data.get('is_staff', False)
        age = request.data.get('age')

        if isinstance(is_superuser, str):
            is_superuser = is_superuser.lower() in ('true', '1', 'yes')
        if isinstance(is_staff, str):
            is_staff = is_staff.lower() in ('true', '1', 'yes')

        if not username or not password:
            return Response(
                {'error': 'Username and password are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        if User.objects.filter(username=username).exists():
            return Response(
                {'error': 'User with this username already exists'},
                status=status.HTTP_400_BAD_REQUEST
            )

        user = User.objects.create_user(
            username=username,
            email=email,
            password=password
        )
        user.is_superuser = is_superuser
        user.is_staff = is_staff
        user.save()
        
        if age:
            try:
                profile = UserProfile.objects.get(user=user)
            except UserProfile.DoesNotExist:
                profile = UserProfile.objects.create(user=user)
            profile.age = int(age)
            profile.save()

        serializer = self.get_serializer(user)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def update(self, request, *args, **kwargs):
        partial = kwargs.pop('partial', False)
        instance = self.get_object()

        if 'username' in request.data:
            instance.username = request.data['username']
        if 'email' in request.data:
            instance.email = request.data['email']
        
        if 'is_superuser' in request.data:
            is_superuser = request.data['is_superuser']
            if isinstance(is_superuser, str):
                is_superuser = is_superuser.lower() in ('true', '1', 'yes')
            instance.is_superuser = is_superuser
            
        if 'is_staff' in request.data:
            is_staff = request.data['is_staff']
            if isinstance(is_staff, str):
                is_staff = is_staff.lower() in ('true', '1', 'yes')
            instance.is_staff = is_staff

        if 'password' in request.data and request.data['password']:
            instance.set_password(request.data['password'])

        instance.save()
        
        if 'age' in request.data:
            age = request.data['age']
            if age:
                try:
                    profile = UserProfile.objects.get(user=instance)
                except UserProfile.DoesNotExist:
                    profile = UserProfile.objects.create(user=instance)
                profile.age = int(age)
                profile.save()

        serializer = self.get_serializer(instance)
        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def stats(self, request):
        queryset_users = self.get_queryset().filter(is_superuser=False)
        total_users = queryset_users.count()

        queryset_admins = self.get_queryset().filter(is_superuser=True)
        total_admins = queryset_admins.count()

        user_profiles = UserProfile.objects.filter(user__in=queryset_users, age__isnull=False)
        if user_profiles.exists():
            age_list = []
            for profile in user_profiles:
                if profile.age:
                    age_list.append(profile.age)
            if age_list:
                avg_age_users = round(sum(age_list) / len(age_list), 1)
            else:
                avg_age_users = 0
        else:
            avg_age_users = 0

        admin_profiles = UserProfile.objects.filter(user__in=queryset_admins, age__isnull=False)
        if admin_profiles.exists():
            admin_age_list = []
            for profile in admin_profiles:
                if profile.age:
                    admin_age_list.append(profile.age)
            if admin_age_list:
                avg_age_admins = round(sum(admin_age_list) / len(admin_age_list), 1)
            else:
                avg_age_admins = 0
        else:
            avg_age_admins = 0

        loans_total = Loan.objects.filter(member__user__in=queryset_users).count()
        
        if total_users > 0:
            avg_books = round(loans_total / total_users, 1)
        else:
            avg_books = 0

        return Response({
            'count_users': total_users,
            'count_admins': total_admins,
            'avg_age_users': avg_age_users,
            'avg_age_admins': avg_age_admins,
            'avg_books': avg_books
        })

    @action(detail=False, methods=['get'], url_path='export/excel')
    def export_excel(self, request):
        queryset = self.get_queryset()
        
        data_list = []
        for user in queryset:
            data_list.append({
                'ID': user.id,
                'Username': user.username,
                'Email': user.email,
                'Is Superuser': 'Yes' if user.is_superuser else 'No',
                'Is Staff': 'Yes' if user.is_staff else 'No'
            })
        
        columns = ['ID', 'Username', 'Email', 'Is Superuser', 'Is Staff']
        
        return self.export_queryset(data_list, columns, 'Members')

    @action(detail=False, methods=['get'], url_path='export/word')
    def export_word(self, request):
        queryset = self.get_queryset()
        
        data_list = []
        for user in queryset:
            data_list.append({
                'ID': user.id,
                'Username': user.username,
                'Email': user.email,
                'Is Superuser': 'Yes' if user.is_superuser else 'No',
                'Is Staff': 'Yes' if user.is_staff else 'No'
            })
        
        columns = ['ID', 'Username', 'Email', 'Is Superuser', 'Is Staff']
        
        from django.http import QueryDict
        request_copy = request._request.GET.copy()
        request_copy['type'] = 'word'
        request._request.GET = request_copy
        
        return self.export_queryset(data_list, columns, 'Members')


class LibraryMemberViewSet(viewsets.ModelViewSet, BaseExportMixin):
    serializer_class = MemberSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = Member.objects.all().select_related('user', 'library')
        
        if self.request.user.is_superuser:
            return queryset
        
        user_member = queryset.filter(user=self.request.user).first()
        
        if not user_member:
            default_library = Library.objects.first()
            
            if default_library:
                user_member = Member.objects.create(
                    user=self.request.user,
                    first_name=self.request.user.username,
                    library=default_library
                )
        
        return queryset.filter(user=self.request.user)
    
    def perform_create(self, serializer):
        if not self.request.user.is_superuser:
            raise PermissionError("Only admins can add library members.")
        serializer.save()
    
    def perform_update(self, serializer):
        if not self.request.user.is_superuser:
            raise PermissionError("Only admins can edit library members.")
        serializer.save()
    
    def perform_destroy(self, instance):
        if not self.request.user.is_superuser:
            raise PermissionError("Only admins can delete library members.")
        instance.delete()

    @action(detail=False, methods=['get'])
    def export(self, request):
        queryset = self.get_queryset()
        
        if not request.user.is_superuser:
            data_list = []
            for member in queryset:
                library_name = member.library.name if member.library else ''
                data_list.append({
                    'ID': member.id,
                    'Library': library_name,
                    'User': '*****'
                })
            columns = ['ID', 'Library', 'User']
        else:
            data_list = []
            for member in queryset:
                member_name = member.first_name
                library_name = member.library.name if member.library else ''
                user_name = member.user.username if member.user else ''
                data_list.append({
                    'ID': member.id,
                    'Name': member_name,
                    'Library': library_name,
                    'User': user_name
                })
            columns = ['ID', 'Name', 'Library', 'User']
        
        return self.export_queryset(data_list, columns, 'LibraryMembers')
